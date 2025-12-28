"""CV processing service for SecureHR application."""

import hashlib
import io
import logging
from pathlib import Path
from typing import Optional, Tuple
from datetime import datetime

import PyPDF2
from docx import Document
from fastapi import HTTPException, UploadFile
from sqlalchemy.orm import Session

from ..database import get_db
from ..models.database import CVVectorDB
from .cyborgdb_service import CyborgDBService

logger = logging.getLogger(__name__)

# File size limit: 10MB
MAX_FILE_SIZE = 10 * 1024 * 1024

# Allowed file extensions
ALLOWED_EXTENSIONS = {".pdf", ".doc", ".docx"}

# MIME types for validation
ALLOWED_MIME_TYPES = {
    "application/pdf",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


class CVProcessorService:
    """Service for handling CV file uploads and text extraction."""

    def __init__(self):
        """Initialize CV processor with CyborgDB service."""
        self.cyborgdb_service = CyborgDBService()

    @staticmethod
    def validate_file(file: UploadFile) -> None:
        """
        Validate uploaded file format, size, and security.
        
        Args:
            file: The uploaded file to validate
            
        Raises:
            HTTPException: If file validation fails
        """
        # Check file size
        if file.size and file.size > MAX_FILE_SIZE:
            raise HTTPException(
                status_code=413,
                detail=f"File size exceeds maximum limit of {MAX_FILE_SIZE // (1024*1024)}MB"
            )
        
        # Check file extension
        if file.filename:
            file_ext = Path(file.filename).suffix.lower()
            if file_ext not in ALLOWED_EXTENSIONS:
                raise HTTPException(
                    status_code=400,
                    detail=f"File type {file_ext} not supported. Allowed types: {', '.join(ALLOWED_EXTENSIONS)}"
                )
        
        # Check MIME type
        if file.content_type not in ALLOWED_MIME_TYPES:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid file type. Content-Type: {file.content_type}"
            )

    @staticmethod
    async def extract_text_from_pdf(file_content: bytes) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_content: PDF file content as bytes
            
        Returns:
            Extracted text content
            
        Raises:
            HTTPException: If PDF processing fails
        """
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            if len(pdf_reader.pages) == 0:
                raise HTTPException(
                    status_code=400,
                    detail="PDF file appears to be empty or corrupted"
                )
            
            text_content = []
            for page in pdf_reader.pages:
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(page_text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from PDF page: {e}")
                    continue
            
            extracted_text = "\n".join(text_content).strip()
            
            if not extracted_text:
                raise HTTPException(
                    status_code=400,
                    detail="No text content could be extracted from PDF"
                )
            
            return extracted_text
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"PDF text extraction failed: {e}")
            raise HTTPException(
                status_code=500,
                detail="Failed to process PDF file"
            )

    @staticmethod
    async def extract_text_from_docx(file_content: bytes) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_content: DOCX file content as bytes
            
        Returns:
            Extracted text content
            
        Raises:
            HTTPException: If DOCX processing fails
        """
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text_content.append(cell.text)
            
            extracted_text = "\n".join(text_content).strip()
            
            if not extracted_text:
                raise HTTPException(
                    status_code=400,
                    detail="No text content could be extracted from DOCX"
                )
            
            return extracted_text
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"DOCX text extraction failed: {e}")
            raise HTTPException(
                status_code=500,
                detail="Failed to process DOCX file"
            )

    @staticmethod
    async def extract_text_from_doc(file_content: bytes) -> str:
        """
        Extract text from DOC file.
        
        Note: This is a simplified implementation. For production use,
        consider using python-docx2txt or antiword for better DOC support.
        
        Args:
            file_content: DOC file content as bytes
            
        Returns:
            Extracted text content
            
        Raises:
            HTTPException: If DOC processing fails
        """
        # For now, we'll raise an error for DOC files as they require
        # additional libraries or system dependencies for proper extraction
        raise HTTPException(
            status_code=400,
            detail="DOC file format not fully supported. Please convert to DOCX or PDF."
        )

    @classmethod
    async def extract_text(cls, file: UploadFile) -> Tuple[str, str]:
        """
        Extract text content from uploaded CV file.
        
        Args:
            file: The uploaded file
            
        Returns:
            Tuple of (extracted_text, file_hash)
            
        Raises:
            HTTPException: If file processing fails
        """
        # Validate file first
        cls.validate_file(file)
        
        # Read file content
        file_content = await file.read()
        
        # Calculate file hash for integrity checking
        file_hash = hashlib.sha256(file_content).hexdigest()
        
        # Reset file position for potential re-reading
        await file.seek(0)
        
        # Extract text based on file type
        file_ext = Path(file.filename or "").suffix.lower()
        
        if file_ext == ".pdf":
            extracted_text = await cls.extract_text_from_pdf(file_content)
        elif file_ext == ".docx":
            extracted_text = await cls.extract_text_from_docx(file_content)
        elif file_ext == ".doc":
            extracted_text = await cls.extract_text_from_doc(file_content)
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file extension: {file_ext}"
            )
        
        # Validate extracted text
        if len(extracted_text.strip()) < 50:
            raise HTTPException(
                status_code=400,
                detail="Extracted text is too short. Please ensure your CV contains sufficient content."
            )
        
        return extracted_text, file_hash

    async def process_cv_complete(
        self, 
        file: UploadFile, 
        candidate_id: str,
        db: Session
    ) -> str:
        """
        Complete CV processing pipeline: extract text and store in CyborgDB.
        
        Args:
            file: The uploaded CV file
            candidate_id: ID of the candidate
            db: Database session
            
        Returns:
            CyborgDB item ID (same as candidate_id)
            
        Raises:
            HTTPException: If any step of processing fails
        """
        try:
            # Step 1: Extract text from file
            extracted_text, file_hash = await self.extract_text(file)
            logger.info(f"Extracted {len(extracted_text)} characters from CV for candidate {candidate_id}")
            
            # Step 2: Store CV text in CyborgDB (it will handle embedding generation automatically)
            metadata = {
                "original_filename": file.filename,
                "file_hash": file_hash,
                "text_length": len(extracted_text),
                "processed_at": datetime.utcnow().isoformat()
            }
            
            cyborgdb_item_id = await self.cyborgdb_service.store_vector(
                cv_text=extracted_text,  # Store original text, CyborgDB handles embeddings
                candidate_id=candidate_id,
                metadata=metadata
            )
            
            # Step 3: Store or update vector metadata in local database (1 user = 1 CV)
            existing_vector = db.query(CVVectorDB).filter(
                CVVectorDB.candidate_id == candidate_id
            ).first()
            
            if existing_vector:
                # Update existing record
                existing_vector.cyborgdb_vector_id = cyborgdb_item_id
                existing_vector.vector_dimensions = "384"  # all-MiniLM-L6-v2 produces 384-dimensional vectors
                existing_vector.original_filename = file.filename
                existing_vector.file_hash = file_hash
            else:
                # Create new record
                vector_record = CVVectorDB(
                    candidate_id=candidate_id,
                    cyborgdb_vector_id=cyborgdb_item_id,
                    vector_dimensions="384",  # all-MiniLM-L6-v2 produces 384-dimensional vectors
                    original_filename=file.filename,
                    file_hash=file_hash
                )
                db.add(vector_record)
            
            db.commit()
            
            logger.info(f"Successfully processed CV for candidate {candidate_id}, CyborgDB ID: {cyborgdb_item_id}")
            return cyborgdb_item_id
            
        except HTTPException:
            raise
        except Exception as e:
            logger.error(f"CV processing failed for candidate {candidate_id}: {e}")
            db.rollback()
            raise HTTPException(
                status_code=500,
                detail=f"CV processing failed: {str(e)}"
            )